import wx
import docx

import os
import mimetypes
import language_tool_python

import threading

import webbrowser


class MyApp(wx.App):
    def OnInit(self):
        self.frame = wx.Frame(None, title="Corrigeo.wx (beta 0.1)", size=(720, 480))
        self.panel = wx.Panel(self.frame)
        
        #LES VARIABLES GLOBALES 
        self.doc_path = None
        self.doc_path_saved = None
        self.doc_info = {}
        self.xdoc = None #le doc à utiliser
        
        # les outils de correction externe
        self.tool = language_tool_python.LanguageTool('fr') 
        self.is_bad_rule = lambda rule: rule.message == 'Faute de frappe possible trouvée.' and len(rule.replacements) and rule.replacements[0][0].isupper()
        # ----------------------------
        
        # Crée un sizer pour le panneau principal
        self.main_sizer = wx.BoxSizer(wx.VERTICAL)
        self.panel.SetSizer(self.main_sizer)

        # PANNEAU AFFICHER POUR AJOUTER UN DOCUMENT
        self.loaddoc_panel = wx.Panel(self.panel)
        loaddoc_panel_size = wx.BoxSizer(wx.VERTICAL)
        self.loaddoc_panel.SetSizer(loaddoc_panel_size)
        
        # PANNEAU DES BOUTONS // ceci est un panneau sous-panneau de loaddoc_panel
        self.btn_panel = wx.Panel(self.loaddoc_panel)
        btn_panel_size = wx.BoxSizer(wx.HORIZONTAL)
        self.btn_panel.SetSizer(btn_panel_size)
        
        
        # LES BOUTONS
        self.bt_opendoc = wx.Button(self.btn_panel, label="Importer un document")
        self.bt_opendoc.Bind(wx.EVT_BUTTON,self.fnOnOpenDoc)
        
        self.bt_correctdoc = wx.Button(self.btn_panel, label="Corriger")
        self.bt_correctdoc.Bind(wx.EVT_BUTTON,self.fnOnCorrectDoc)
        self.bt_correctdoc.Enable(False)
        
        self.bt_savedoc = wx.Button(self.btn_panel, label="Sauvegarder le fichier")
        self.bt_savedoc.Bind(wx.EVT_BUTTON,self.fnOnSaveDoc)
        self.bt_savedoc.Enable(False)
        self.bt_opensaveddoc = wx.Button(self.btn_panel, label="Ouvrir")
        self.bt_opensaveddoc.Bind(wx.EVT_BUTTON,self.fnOnLaunchDocxFile)
        self.bt_opensaveddoc.Enable(False)

        self.lb_info = wx.StaticText(self.loaddoc_panel, label="Vous pouvez importer un document Word ici")
        
        self.progress_bar = wx.Gauge(self.loaddoc_panel, range=100)
        # self.progress_bar.Hide()
        
        btn_panel_size.Add(self.bt_opendoc,0,wx.ALIGN_CENTER)
        btn_panel_size.Add(self.bt_correctdoc,0,wx.ALIGN_CENTER)
        btn_panel_size.AddSpacer(20)
        
        btn_panel_size.Add(self.bt_savedoc,0,wx.ALIGN_CENTER)
        btn_panel_size.Add(self.bt_opensaveddoc,0,wx.ALIGN_CENTER)
        
        loaddoc_panel_size.AddStretchSpacer(1)
        loaddoc_panel_size.Add(self.lb_info,0,wx.ALIGN_CENTER)
        loaddoc_panel_size.AddSpacer(20)
        loaddoc_panel_size.Add(self.btn_panel, 0, wx.ALIGN_CENTER)
        loaddoc_panel_size.AddSpacer(20)
        loaddoc_panel_size.Add(self.progress_bar, 0, wx.ALIGN_CENTER)
        loaddoc_panel_size.AddStretchSpacer(1)
        
        
        # Ajoute le panneau contenant le bouton au sizer principal
        self.main_sizer.Add(self.loaddoc_panel, 1, wx.EXPAND | wx.ALL, 5)

        self.frame.Show(True)
        return True
    
    def fnOnOpenDoc(self, event):
        wildcard = "Fichiers Word (*.docx)|*.docx"  # Définit les types de fichiers acceptés
        dialog = wx.FileDialog(self.frame, "Sélectionnez un fichier .docx", wildcard=wildcard, style=wx.FD_OPEN)

        if dialog.ShowModal() == wx.ID_OK:
            selected_file = dialog.GetPath()
            self.doc_path = selected_file
            self.fnGetDocInfo(self.doc_path)
            print("Fichier sélectionné : ", selected_file)

        dialog.Destroy()
    
    def fnOnSaveDoc(self, event):
        wildcard = "Fichiers Word (*.docx)|*.docx"  # Filtre pour les fichiers .docx
        file_dialog = wx.FileDialog(self.frame, message="Enregistrer un fichier .docx", wildcard=wildcard, style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT)

        # Définir le nom de fichier par défaut
        nom = os.path.splitext(self.doc_info['nom'])[0]
        file_dialog.SetFilename(f"{nom}-corr.docx")
        
        if file_dialog.ShowModal() == wx.ID_CANCEL:
            file_dialog.Destroy()
            return

        file_path = file_dialog.GetPath()
        print(file_path)
        self.doc_path_saved = file_path
        
        self.bt_opensaveddoc.Enable(True)
        self.xdoc.save(file_path)
        
        file_dialog.Destroy()
    
    def fnOnLaunchDocxFile(self, event):
        webbrowser.open(self.doc_path_saved)
    
    # QUE DES GESTIONS D'ETATS DE BOUTONS
    def fnDisableAllButton(self):
        self.bt_correctdoc.Enable(False)
        self.bt_opendoc.Enable(False)
        self.bt_opensaveddoc.Enable(False)
        self.bt_savedoc.Enable(False)
    
    def fnEnableCorrectButton(self):
        self.bt_correctdoc.Enable(True)
        self.bt_opendoc.Enable(True)
        self.bt_opensaveddoc.Enable(False)
        self.bt_savedoc.Enable(False)
        
    def fnEnableSaveDocButton(self):
        self.bt_correctdoc.Enable(False)
        self.bt_opendoc.Enable(True)
        self.bt_opensaveddoc.Enable(False)
        self.bt_savedoc.Enable(True)
    ##############################
        
    def fnOnCorrectDoc(self, event):
        # bah ici la correction de doc 
        self.fnDisableAllButton()
        # self.progress_bar.Show()
        
        correction_thread = threading.Thread(target=self.fnCorrectionDoc)
        correction_thread.start()
        
    def fnUpdateProgress(self,value):
        self.progress_bar.SetValue(value)
    
    def fnCorrectionDoc(self):
        
        run_separator = "[Run]"
        nb_paragraph = len(self.xdoc.paragraphs)
        
        for index, p in enumerate(self.xdoc.paragraphs):
            text = ""
            print(f"{index+1}/{nb_paragraph}")
            prc = (index +1) * 100 / nb_paragraph
            wx.CallAfter(self.fnUpdateProgress, int(prc))
            print(prc)
            # Suppression de run
            last_run = 0
            runs = p.runs
            
            index_runs_to_delete = []
            run_delete = 0
            
            # print(len(p.runs))
            # cette booucle permet de regroupé les runs afin de facilité la correction
            for index,run in enumerate(runs):
                if index > 0 :
                    if run.bold == p.runs[last_run].bold and run.italic == p.runs[last_run].italic and run.underline == p.runs[last_run].underline and run.font.color.rgb == p.runs[last_run].font.color.rgb and run.font.size == p.runs[last_run].font.size and run.font.underline == p.runs[last_run].font.underline:
                        p.runs[last_run].text += run.text
                        run_index = index - run_delete
                        p.runs[run_index]._element.getparent().remove(p.runs[run_index]._element)
                        run_delete += 1
                    else:
                        last_run = index - run_delete
                
            # print(len(p.runs))
            
            for index,run in enumerate(p.runs):
                if index != 0:
                    text += run_separator
                text += run.text
                # print(f"__run__ : {run.text}")
                
            # print("___req___")
            if len(p.runs) <= 0:
                text = p.text
                
            matches = self.tool.check(text)
            matches = [rule for rule in matches if not self.is_bad_rule(rule)]
            # print(matches)
            corrrige = language_tool_python.utils.correct(text, matches)
            # print(corrrige)
            
            if len(p.runs) <= 0:
                p.text = corrrige
            else:
                res_split = corrrige.split(run_separator)
                # print(f"Taille : {len(res_split)}")
                for index,tt in enumerate(res_split):
                    p.runs[index].text = res_split[index]
            
            # print("___ end req ___")
    
        self.fnEnableSaveDocButton()
        
        default_font = wx.SystemSettings.GetFont(wx.SYS_DEFAULT_GUI_FONT)
        self.lb_info.SetFont(default_font)
        self.lb_info.SetLabelText("La corréction s'est bien terminé. Vous pouvez sauvegarder le fichier et l'ouvrir.")
        self.loaddoc_panel.Layout()

    def fnGetDocInfo(self,doc_path):
        # Obtenir la taille du fichier (en octets)
        taille_du_fichier = os.path.getsize(doc_path)

        # Obtenir le nom du fichier (avec extension)
        nom_du_fichier = os.path.basename(doc_path)

        # Obtenir l'extension du fichier
        extension_du_fichier = os.path.splitext(nom_du_fichier)[1]
        
        mime_type, _ = mimetypes.guess_type(doc_path)
        est_document_word = mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        res = {
            'taille':taille_du_fichier, #en octets
            'nom':nom_du_fichier,
            'extension':extension_du_fichier,
            'isWord':est_document_word
        }
        
        self.doc_info = res
        
        if est_document_word:
            self.xdoc = docx.Document(doc_path)
            self.lb_info.SetLabelText(f"Le fichier '{res['nom']}' est valide.")
            
            self.lb_info.SetForegroundColour(wx.Colour(10, 200, 0))  # Rouge

            # Change la taille de la police
            font = wx.Font(12, wx.DEFAULT, wx.NORMAL, wx.NORMAL)
            self.lb_info.SetFont(font)
            
            self.fnEnableCorrectButton()
            
            self.bt_correctdoc.SetFocus()
            self.loaddoc_panel.Layout()
        else:
            self.lb_info.SetLabelText(f"Le fichier '{res['nom']}' n'est pas un document word valide.")
            
            self.lb_info.SetForegroundColour(wx.Colour(255, 0, 0))  # Rouge

            # Change la taille de la police
            font = wx.Font(16, wx.DEFAULT, wx.NORMAL, wx.NORMAL)
            self.lb_info.SetFont(font)
            
            self.bt_correctdoc.Enable(False)
            self.bt_opendoc.SetFocus()
            
            self.loaddoc_panel.Layout()
            
if __name__ == '__main__':
    app = MyApp()
    app.MainLoop()
